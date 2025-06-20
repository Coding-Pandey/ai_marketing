from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Body, Form, Query
from fastapi.responses import JSONResponse
from typing import Optional
from social_media.Agents.document_summared import Document_summerizer
from S3_bucket.fetch_document import download_document
from S3_bucket.fetch_document import download_document
from content_generation.blog_agent.blog_generation import blog_generation
from content_generation.blog_agent.blog_suggest import blog_suggest
from content_generation.blog_agent.seo_blog import generation_blog_async
from utils import verify_jwt_token, check_api_limit
import json
from auth.auth import get_db
from sqlalchemy.orm.attributes import flag_modified
from sqlalchemy.orm import Session
from auth.models import Contentgeneration, ContentgenerationFile, SourceFileContent
from datetime import timedelta
from content_generation.content_generation_model import UUIDRequest,ContentGenerationFileSchema
import uuid
import os
from docx import Document
import shutil


router = APIRouter()

TEMP_FILE_DIR = "content_generation/tmp/uploads"  # Ensure this directory exists

content_types = [
    {"id": 1, "content_type": "blog generation"},
    {"id": 2, "content_type": "information post"},
    {"id": 3, "content_type": "press release and news pages"},
    {"id": 4, "content_type": "case study pages"},
    {"id": 5, "content_type": "FAQs"},
    {"id": 6, "content_type": "campaign landing pages"}, 
    {"id": 7, "content_type": "product descriptions"}  
]

@router.get("/content_types")
async def get_content_types():
    try:
        return {"content_types": content_types}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch content types: {str(e)}")


@router.post("/content_generation")
async def content_generation(
    file: Optional[UploadFile] = File(None),
    text_data: Optional[str] = Form(None),
    content_type: Optional[int] = Form(None), 
    file_name: Optional[str] = Form(None),   
    objectives: Optional[str] = Form(None),
    audience: Optional[str] = Form(None),
    keywords: Optional[str] = Form(None),
    user=Depends(check_api_limit("content_generation")),
    db: Session = Depends(get_db),
    user_id: str = Depends(verify_jwt_token)
):
    try:
        user_id = int(user_id[1])  # Extract actual user ID from JWT payload

        # Validate required fields
        if not file_name:
            raise HTTPException(status_code=400, detail="File name is required")

        if not file and not text_data:
            raise HTTPException(status_code=400, detail="Either file or text_data must be provided")

        # Validate file format
        if file and not file.filename.lower().endswith((".docx", ".doc", ".pdf")):
            raise HTTPException(status_code=400, detail="Invalid file format. Please upload a .docx, .doc, or .pdf file")

        if content_type != 1:    
            return JSONResponse(
                status_code=400,
                content={"message": "Only blog generation is supported at this time"}
            )
        
        # Validate content type
        allowed_content_types = [1]  # Expand this if needed
        if content_type and content_type not in allowed_content_types:
            raise HTTPException(status_code=400, detail="Invalid content type")
        
        # Process keywords - handle both empty/None and actual keyword data
        processed_keywords = None
        if keywords:
            try:
                keywords_data = json.loads(keywords)
                # Extract keywords from the nested structure
                if "Keywords" in keywords_data and isinstance(keywords_data["Keywords"], list):
                    processed_keywords = keywords_data
                    file_name = keywords_data.get("Page_Title", "")
                else:
                    # Handle old format for backward compatibility
                    data = keywords_data.get("keywords", [])
                    file_name = keywords_data.get("Page_Title", "")
                    if data:
                        processed_keywords = [keyword.strip() for keyword in data if isinstance(keyword, str)]
            except json.JSONDecodeError:
                raise HTTPException(status_code=400, detail="Invalid JSON format for keywords")
            
        

        # Handle file or text content
        if file:
            file_contents = await file.read()

            # Save file temporarily
            TEMP_FILE_DIR = "content_generation/tmp/uploads/User_" + str(user_id)
            # remove the entire directory if it exists
            if os.path.isdir(TEMP_FILE_DIR):
                shutil.rmtree(TEMP_FILE_DIR)
            
            os.makedirs(TEMP_FILE_DIR, exist_ok=True)
            temp_file_path = os.path.join(TEMP_FILE_DIR, f"{uuid.uuid4().hex}_{file.filename}")
            with open(temp_file_path, "wb") as f:
                f.write(file_contents)

            temp_file_path = os.path.normpath(temp_file_path).replace("\\", "/")    

        else:
            file_contents = text_data.encode("utf-8")
            temp_file_path = None  # No physical file saved

        if objectives:
            objectives = json.loads(objectives)

        if audience:
            audience = json.loads(audience)

        filecontent_obj = []
        filecontent = db.query(SourceFileContent).filter(SourceFileContent.user_id == user_id).all()
        if filecontent is not None:
            for obj in objectives:
                # print(obj)
                for i in filecontent:
                    # print(i.uuid_id)
                    if i.uuid_id == obj:
                        fileData = i.file_data
                        # print(fileData)
                    # filecontent = i.extracted_text
                        filecontent_obj.append(fileData)

            for obj in audience:
                for i in filecontent:
                    if i.uuid_id == obj:
                        audienceData = i.file_data
                    # filecontent = i.extracted_text
                        filecontent_obj.append(audienceData)
        if filecontent_obj:
            summarized_text_json, total_toke = Document_summerizer(filecontent_obj)
        else:
            summarized_text_json = {}
        
        print(summarized_text_json)    

        # summarized_text_json = {}
        total_tokens = 0

        if content_type == 1:
            try:
                json_data, total_tokens = blog_generation(
                    file=file_contents,
                    json_data=summarized_text_json,
                    keywords=processed_keywords,
                )

                # Update user token usage
                content_gen_record = db.query(Contentgeneration).filter(Contentgeneration.user_id == user_id).first()
                if content_gen_record:
                    if total_tokens > 0:
                        content_gen_record.total_tokens += total_tokens
                    if not json_data:
                        content_gen_record.call_count = max(content_gen_record.call_count - 1, 0)
                    db.commit()
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Blog generation failed: {str(e)}")
            

        return JSONResponse(content={
            "filename": file_name,
            "content_type": "blog generation",
            "data": json_data,
            "temp_file_path": temp_file_path  # Send this if you want to reuse the file
        })
    

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Content generation failed: {str(e)}")


# @router.post("/content_generation_new")


@router.get("/content_datalist")
async def content_documents(db: Session = Depends(get_db), id: str = Depends(verify_jwt_token)):
    try:
        user_id = int(id[1])  
        content_files = db.query(ContentgenerationFile).filter(ContentgenerationFile.user_id == user_id).all()
        if not content_files:
            return []

        file_count = len(content_files)


        content_record = db.query(Contentgeneration).filter(Contentgeneration.user_id == user_id).first()

        if content_record:

            content_record.file_count = file_count
            content_record.call_count = file_count
            db.commit()

   
        result = [
            {
                "file_name": content_file.file_name,
                "uuid": content_file.uuid,
                "last_reset": content_file.last_reset + timedelta(days=30) if content_file.last_reset else None,
            }
            for content_file in content_files
        ]

        return result

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))    

@router.delete("/content_delete_data")
async def content_delete_data(request: UUIDRequest, id: str = Depends(verify_jwt_token), db: Session = Depends(get_db)):
    try:
        user_id = str(id[1])  # Extract user_id from the JWT token
        uuid = request.uuid
        # success = seo_cluster_delete_document(uuid, user_id)
        # if success:
            # 2. Delete DB record
        user_id = int(id[1]) 
        file_record = db.query(ContentgenerationFile).filter_by(user_id=user_id, uuid=uuid).first()

        if file_record:
            db.delete(file_record)
            db.commit()
            # db.refresh(file_record) 

        return JSONResponse(
            status_code=200,
            content={
                "message": "Document deleted successfully",
                "uuid": uuid
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))   
    except Exception as e:
        db.rollback()
        raise Exception(f"Error storing file in table: {str(e)}")
    finally:
        db.close()  

@router.get("/content_data/{uuid}")
async def content_fetch_data(uuid: str, id: str = Depends(verify_jwt_token), db: Session = Depends(get_db)):
    try:
        if uuid is None:
            raise HTTPException(status_code=200, detail="UUID is required")
        user_id = str(id[1]) 
        print(user_id) # Extract user_id from the JWT token
        uuid = str(uuid)  # Ensure uuid is a string
        file = db.query(ContentgenerationFile).filter_by(user_id=user_id, uuid=uuid).first()

        # print(f"json_data: {seo_file.json_data}")
        if not file:
            raise HTTPException(status_code=200, detail="File not found for this user")
        
        # data = fetch_seo_cluster_file(user_id, uuid)
        # if not data:
        #     raise HTTPException(status_code=200, detail="No documents found for the user")
        
        # json_data = json.loads(data["documents"])
        

        json_data = {
            "id": uuid,
            "filename": file.file_name,
            "content_type": file.content_type,
            "data":  file.content_data
        }
        return json_data

    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.patch("/content_generation_uploadfile/{uuid_id}")
async def content_generation_upload_file(
    json_data: ContentGenerationFileSchema = Body(...),
    uuid_id: Optional[str] = None,
    id: str = Depends(verify_jwt_token),
    temp_file_path: Optional[str] = None,
    db: Session = Depends(get_db)
):
    try:
        # Extract user_id from JWT token
        user_id = id[1]  # Assuming it's a tuple like (status, user_id)

        if uuid_id != "undefined":
            db_query = db.query(ContentgenerationFile).filter(ContentgenerationFile.uuid == uuid_id)
            content_generation = db_query.first()

            if not content_generation:
                raise HTTPException(status_code=404, detail="Record not found for provided UUID.")

            # Update existing entry
            content_generation.content_data = json_data.data
            content_generation.file_name = json_data.filename
            
            flag_modified(content_generation, "content_data")
            db.commit()
            db.refresh(content_generation)

            return JSONResponse(
                status_code=200,
                content={"message": "File updated successfully", "uuid": uuid_id}
            )

        else:
            # Validation
            file_content = json_data.data
            file_name = json_data.filename
            content_type = json_data.content_type

            if not file_content:
                raise HTTPException(status_code=400, detail="No data provided")
            if not file_name:
                raise HTTPException(status_code=400, detail="No filename provided")
            if not content_type:
                raise HTTPException(status_code=400, detail="No content type provided")

            # Create new entry
            unique_id = uuid.uuid4().hex

            content_generation = ContentgenerationFile(
                user_id=user_id,
                file_name=file_name,
                uuid=unique_id,
                content_type=content_type,
                content_data=file_content
            )

            db.add(content_generation)
            flag_modified(content_generation, "content_data")
            db.commit()
            db.refresh(content_generation)

            return JSONResponse(
                status_code=200,
                content={
                    "message": "File uploaded successfully",
                    "uuid": unique_id,
                    "filename": file_name,
                    "content_type": content_type
                }
            )

    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred: {str(e)}"
        )
    finally:
        # Safe temp file deletion
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except Exception as cleanup_error:
                # Optional: log cleanup failure
                print(f"Failed to remove temp file: {cleanup_error}")

        else:
            print("No temp file to delete or file does not exist.")        


@router.post("/blog_suggestion_more")
async def blog_suggestion_more(
    file_path: Optional[str] = Form(None),
    text_data: Optional[str] = Form(None),
    generated_blog: Optional[str] = Form(None),
    # user = Depends(check_api_limit("content_generation")),
    db: Session = Depends(get_db),
    user_id: str = Depends(verify_jwt_token)   # Renamed to be more descriptive
):
    try:
        user_id = int(user_id[1])  # Extract user_id from the JWT token
        # Validate inputs

        if not file_path and not text_data:
            raise HTTPException(status_code=400, detail="Either file or text_data must be provided")

        # # Validate file format
        # if file_path and not file_path.filename.lower().endswith((".docx", ".doc", ".pdf")):
        #     raise HTTPException(status_code=400, detail="Invalid file format. Please upload a .docx, .doc, or .pdf file")

        # Process file or text data
        if file_path:
            # with open(file_path, "r", encoding="utf-8") as f:
            #     content = f.read()
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            print(f"Content from file: {content[:100]}...")  # Log first 100 characters for debugging

        file_contents = content if file_path else text_data.encode('utf-8')

        # Initialize variables
        summarized_text_json = {}
        total_token = 0
        try:
            # Assuming blog_generation returns JSON data and token count
            json_data, total_tokens = blog_suggest(
                file=file_contents,
                Generated_Blog=generated_blog,
                json_data=summarized_text_json
            )
            total_token += total_tokens

            return json_data

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Blog generation failed: {str(e)}")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Content generation failed: {str(e)}")


@router.patch("/update_name_and_title/{uuid}")
async def update_name_and_title(
    uuid: str,
    new_filename: Optional[str] = Body(None),
    new_title: Optional[str] = Body(None),
    id: str = Depends(verify_jwt_token),
    db: Session = Depends(get_db)
):
    try:
        user_id = int(id[1])

        content_file = db.query(ContentgenerationFile).filter(
            ContentgenerationFile.uuid == uuid,
            ContentgenerationFile.user_id == user_id
        ).first()

        if not content_file:
            raise HTTPException(status_code=404, detail="File not found")

        # If neither field is sent, raise error
        if not new_filename and not new_title:
            raise HTTPException(status_code=400, detail="At least one of new_filename or new_title must be provided")

        if new_filename and new_filename.strip():
            content_file.file_name = new_filename

        if new_title and new_title.strip():
            if isinstance(content_file.content_data, dict):
                content_file.content_data["Title"] = new_title
                flag_modified(content_file, "content_data")
            else:
                raise HTTPException(status_code=500, detail="content_data is not a valid JSON object")

        db.commit()
        db.refresh(content_file)

        return {
            "message": "Update successful",
            "new_filename": content_file.file_name,
            "new_title": content_file.content_data.get("Title", None)
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}")

@router.delete("/delete_temp_file")
def delete_temp_file(file_path: str = Query(..., description="Relative file path from frontend")):
    # Security check: prevent directory traversal
    if ".." in file_path or file_path.startswith("/"):
        raise HTTPException(status_code=400, detail="Invalid file path.")

    # Resolve safe base dir
    base_dir = os.path.abspath("content_generation/tmp/uploads")
    full_path = os.path.abspath(file_path)

    # Ensure it's inside the uploads folder
    if not full_path.startswith(base_dir):
        raise HTTPException(status_code=403, detail="Unauthorized file path")

    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="File not found")

    try:
        os.remove(full_path)
        return {"message": f"File '{file_path}' deleted successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting file: {str(e)}")
    

@router.post("/edit_content_generation")
async def edit_content_generation(
    file: Optional[UploadFile] = File(None),
    temp_file_path: Optional[str] = Form(None),
    text_data: Optional[str] = Form(None),
    content_type: Optional[int] = Form(None), 
    file_name: Optional[str] = Form(None),   
    objectives: Optional[str] = Form(None),
    audience: Optional[str] = Form(None),
    user=Depends(check_api_limit("content_generation")),
    db: Session = Depends(get_db),
    user_id: str = Depends(verify_jwt_token)
):
    try:
        user_id = int(user_id[1])  

        # Validate required inputs
        if not file and not temp_file_path and not text_data:
            raise HTTPException(status_code=400, detail="Provide either a file, a temp_file_path, or text_data")

        if not file_name:
            raise HTTPException(status_code=400, detail="File name is required")

        file_contents = None
        temp_file = None
        # Use uploaded file if available
        if file and temp_file_path:
            # Validate file format
            if not file.filename.lower().endswith((".docx", ".doc", ".pdf")):
                raise HTTPException(status_code=400, detail="Invalid file format. Please upload a .docx, .doc, or .pdf file")
            file_contents = await file.read()

            TEMP_FILE_DIR = "content_generation/tmp/uploads/User_" + str(user_id)

            if os.path.isdir(TEMP_FILE_DIR):
                shutil.rmtree(TEMP_FILE_DIR)

            os.makedirs(TEMP_FILE_DIR, exist_ok=True)
            new_file_path = os.path.join(TEMP_FILE_DIR, f"{uuid.uuid4().hex}_{file.filename}")
            with open(new_file_path, "wb") as f:
                f.write(file_contents)

            temp_file = os.path.normpath(new_file_path).replace("\\", "/")

            # base_dir = os.path.abspath("content_generation/tmp/uploads")
            # full_path = os.path.abspath(temp_file_path)

            # # Ensure it's inside the uploads folder
            # if not full_path.startswith(base_dir):
            #     raise HTTPException(status_code=403, detail="Unauthorized file path")

            # if not os.path.exists(full_path):
            #     raise HTTPException(status_code=404, detail="File not found")
            
            # os.remove(full_path)



        # If no file uploaded, try reading from temp_file_path
        elif temp_file_path and not file:
            try:
                doc = Document(temp_file_path)
                file_contents = "\n".join([para.text for para in doc.paragraphs])
                temp_file = temp_file_path
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to read from temp file path: {str(e)}")

        # Validate content type
        allowed_content_types = [1]  # Expand this if needed
        if content_type and content_type not in allowed_content_types:
            raise HTTPException(status_code=400, detail="Invalid content type")

        summarized_text_json = {}
        total_tokens = 0

        if content_type == 1:
            try:
                json_data, total_tokens = blog_generation(
                    file=file_contents,
                    json_data=summarized_text_json
                )

                # Update user token usage
                content_gen_record = db.query(Contentgeneration).filter(Contentgeneration.user_id == user_id).first()
                if content_gen_record:
                    if total_tokens > 0:
                        content_gen_record.total_tokens += total_tokens
                    if not json_data:
                        content_gen_record.call_count = max(content_gen_record.call_count - 1, 0)
                    db.commit()
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Blog generation failed: {str(e)}")
            
        if content_type != 1:    
            return JSONResponse(
                status_code=400,
                content={"message": "Only blog generation is supported at this time"}
            )

        return JSONResponse(content={
            "filename": file_name,
            "content_type": "blog generation",
            "data": json_data,
            "temp_file_path": temp_file  
        })
    

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Content generation failed: {str(e)}")
    

# @router.post("/seo_based_blog")
# async def Seo_based_blog(csv_data: str = Form(...), text: Optional[str] = Form(None)):
#     try:
#         parsed_data = json.loads(csv_data) 
#         # print(parsed_data)
#         json_data = download_csv(parsed_data['data'])
#         # print(json_data)
#         keywords = [item['keyword'] for item in json_data['csv']]
#         new_blog ,total_tokens_used = await generation_blog_async(keywords, text)
#         print(f"total Seo blog {total_tokens_used}")
#         return new_blog, total_tokens_used

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))    

