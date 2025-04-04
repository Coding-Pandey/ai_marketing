import streamlit as st
import requests
import pandas as pd
import json
from io import BytesIO
from typing import List, Optional
from docx import Document
import io
import os
import subprocess
import tempfile

st.set_page_config(layout="wide")

# Streamlit UI
st.title("AI marketing tool")


# FastAPI backend URLs
GENERATE_API_URL = "http://127.0.0.1:8000/seo_generate_keywords"
SUGGEST_API_URL = "http://127.0.0.1:8000/seo_keyword_suggestion"
CLUSTER_API_URL = "http://127.0.0.1:8000/seo_keyword_clustering"

PPC_GENERATE_API_URL = "http://127.0.0.1:8000/ppc_generate_keywords"
PPC_CLUSTER_API_URL = "http://127.0.0.1:8000/ppc_keyword_clustering"

SOCIAL_MEDIA_API_URL = "http://127.0.0.1:8000/social_media_post"
UPLOAD_FILE_S3_BUCKET = "http://127.0.0.1:8000/uploadfile"

FETCH_DOCUMENTS_URL = "http://127.0.0.1:8000/list-documents"
DOWNLOAD_DOCUMENT_URL = "http://127.0.0.1:8000/process-documents"

SEO_PROCESS_FILE ="http://127.0.0.1:8000/seo_uploadfile"
DOWNLOAD_CONTENT_CSV ="http://127.0.0.1:8000/process_content_generation"

BLOG_GENERATION_TEXT ="http://127.0.0.1:8000/blog_generation"
# Create tabs for SEO and PPC processes
tab1, tab2, tab3, tab4, tab6, tab5 = st.tabs(["SEO Process", 
                                        "PPC Process", 
                                        "Keywords suggestions",
                                        "Social Media Post", 
                                        "Content generation",
                                        "Upload Document",
                                        ])

# Initialize session state for both tabs
if "seo_df" not in st.session_state:
    st.session_state.seo_df = None
if "seo_processed_df" not in st.session_state:
    st.session_state.seo_processed_df = None
if "ppc_df" not in st.session_state:
    st.session_state.ppc_df = None
if "ppc_processed_df" not in st.session_state:
    st.session_state.ppc_processed_df = None

# Sample location and language data (replace with your actual data)
location_options =[
    {"id": 2840, "country": "United States"},
    {"id": 2826, "country": "United Kingdom"},
    {"id": 2616, "country": "Poland"},
    {"id": 2276, "country": "Germany"},
    {"id": 2250, "country": "France"},
    {"id": 2380, "country": "Italy"},
    {"id": 2724, "country": "Spain"},
    {"id": 2124, "country": "Canada"},
    {"id": 2036, "country": "Australia"},
    {"id": 2356, "country": "India"},
    {"id": 2392, "country": "Japan"},
    {"id": 2484, "country": "China"},
    {"id": 2076, "country": "Brazil"},
    {"id": 2156, "country": "Mexico"},
    {"id": 2252, "country": "Russia"},
    {"id": 2190, "country": "South Africa"}
]

language_options = [
    {
        "ID": 1000,
        "Name": "English"
    },
    {
        "ID": 1001,
        "Name": "German"
    },
    {
        "ID": 1002,
        "Name": "French"
    },
    {
        "ID": 1003,
        "Name": "Spanish"
    },
    {
        "ID": 1004,
        "Name": "Italian"
    },
    {
        "ID": 1005,
        "Name": "Japanese"
    },
    {
        "ID": 1009,
        "Name": "Danish"
    },
    {
        "ID": 1010,
        "Name": "Dutch"
    },
    {
        "ID": 1011,
        "Name": "Finnish"
    },
    {
        "ID": 1012,
        "Name": "Korean"
    },
    {
        "ID": 1013,
        "Name": "Norwegian"
    },
    {
        "ID": 1014,
        "Name": "Portuguese"
    },
    {
        "ID": 1015,
        "Name": "Swedish"
    },
    {
        "ID": 1017,
        "Name": "Chinese (simplified)"
    },
    {
        "ID": 1018,
        "Name": "Chinese (traditional)"
    },
    {
        "ID": 1019,
        "Name": "Arabic"
    },
    {
        "ID": 1020,
        "Name": "Bulgarian"
    },
    {
        "ID": 1021,
        "Name": "Czech"
    },
    {
        "ID": 1022,
        "Name": "Greek"
    },
    {
        "ID": 1023,
        "Name": "Hindi"
    },
    {
        "ID": 1024,
        "Name": "Hungarian"
    },
    {
        "ID": 1025,
        "Name": "Indonesian"
    },
    {
        "ID": 1026,
        "Name": "Icelandic"
    },
    {
        "ID": 1027,
        "Name": "Hebrew"
    },
    {
        "ID": 1028,
        "Name": "Latvian"
    },
    {
        "ID": 1029,
        "Name": "Lithuanian"
    },
    {
        "ID": 1030,
        "Name": "Polish"
    },
    {
        "ID": 1031,
        "Name": "Russian"
    },
    {
        "ID": 1032,
        "Name": "Romanian"
    },
    {
        "ID": 1033,
        "Name": "Slovak"
    },
    {
        "ID": 1034,
        "Name": "Slovenian"
    },
    {
        "ID": 1035,
        "Name": "Serbian"
    },
    {
        "ID": 1036,
        "Name": "Ukrainian"
    },
    {
        "ID": 1037,
        "Name": "Turkish"
    },
    {
        "ID": 1038,
        "Name": "Catalan"
    },
    {
        "ID": 1039,
        "Name": "Croatian"
    },
    {
        "ID": 1040,
        "Name": "Vietnamese"
    },
    {
        "ID": 1041,
        "Name": "Urdu"
    },
    {
        "ID": 1042,
        "Name": "Filipino"
    },
    {
        "ID": 1043,
        "Name": "Estonian"
    },
    {
        "ID": 1044,
        "Name": "Thai"
    },
    {
        "ID": 1056,
        "Name": "Bengali"
    },
    {
        "ID": 1064,
        "Name": "Persian"
    },
    {
        "ID": 1072,
        "Name": "Gujarati"
    },
    {
        "ID": 1086,
        "Name": "Kannada"
    },
    {
        "ID": 1098,
        "Name": "Malayalam"
    },
    {
        "ID": 1101,
        "Name": "Marathi"
    },
    {
        "ID": 1102,
        "Name": "Malay"
    },
    {
        "ID": 1110,
        "Name": "Punjabi"
    },
    {
        "ID": 1130,
        "Name": "Tamil"
    },
    {
        "ID": 1131,
        "Name": "Telugu"
    }
]

# SEO Tab
with tab1:
    st.subheader("SEO Keyword Generator & Suggestion")
    # User input field
    seo_keywords = st.text_input(
        "Enter SEO Keywords (comma-separated):", 
        key="seo_key",
        placeholder="e.g., SEO, Machine Learning, Automation"
    )

    # Process input when user submits
    if seo_keywords:
        keywords_list = [kw.strip() for kw in seo_keywords.split(",") if kw.strip()]     
        keywords_list = keywords_list[:10]

        st.success(f"✅ Using the first {len(keywords_list)} keywords: {', '.join(keywords_list)}")


    seo_description = st.text_area("Enter a Short Description (Optional):", key="seo_desc", 
                                placeholder="Describe your business or service for SEO")
    
    # Location and Language dropdowns
    st.subheader("Search Parameters")
    col1, col2 = st.columns(2)
    
    with col1:
        # Multi-select for locations
        selected_locations = st.multiselect(
            "Select Locations",
            options=[loc["country"] for loc in location_options],
            default=[location_options[0]["country"]],
            key="seo_locations"
        )
        
        # Get the location IDs based on selected location names
        location_ids = [
            loc["id"] for loc in location_options 
            if loc["country"] in selected_locations
        ]
    
    with col2:
        # Single select for language
        selected_language = st.selectbox(
            "Select Language",
            options=[lang["Name"] for lang in language_options],
            index=0,
            key="seo_language"
        )
        
        # Get the language ID based on selected language name
        language_id = next(
            (lang["ID"] for lang in language_options if lang["Name"] == selected_language),
            None
        )
    
    # Checkbox to enable/disable the exclusion filter
    use_exclude_filter = st.checkbox("Enable Exclusion Filter")
    if use_exclude_filter:
        max_exclude = st.slider(
            "Exclude Keywords with Monthly Searches Below",
            min_value=0,
            max_value=500,
            value=0,
            step=10,
            key="seo_exclude"
        )
    else:
        max_exclude = None


    # Generate exclusion list
    if max_exclude is None:
        exclude_values = []
    else:
        exclude_values = list(range(0, max_exclude + 1, 10))  # Include max_exclude in the range

        # Ensure the max_exclude is in the list if it's not already there
        if max_exclude > 0 and max_exclude not in exclude_values:
            exclude_values.append(int(max_exclude))



    st.caption(f"Will exclude keywords with monthly searches in: {exclude_values}")

    remove_branded_status = st.checkbox("remove branded keywords")
    if remove_branded_status:
        remove_branded = True
    else:
        remove_branded = False


    if 'seo_word_list' not in st.session_state:
        st.session_state.seo_word_list = []

    seo_new_word = st.text_input("Enter branded Keywords (comma-separated):", key="seo_branded_key", 
                                    placeholder="e.g., apple, samsung, Automation")


    if st.button("Add Keywords"):
        if seo_new_word:
            words = [word.strip() for word in seo_new_word.split(',')]
    
            for word in words:
                if word and word not in st.session_state.seo_word_list:
                    st.session_state.seo_word_list.append(word)
            
      
    show_keywords = st.checkbox("Show Keywords", value=False)

    if show_keywords:
        st.subheader("Current Word List:")
        if st.session_state.seo_word_list:
            # Create columns for each word to add a remove button
            for i, word in enumerate(st.session_state.seo_word_list):
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.write(f"{i+1}. {word}")
                with col2:
                    # Create a unique key for each button
                    if st.button("Remove", key=f"remove_{i}"):
                        st.session_state.seo_word_list.pop(i)
                        st.rerun()  # Rerun the app to update the display
        else:
            st.write("No words added yet.")
    
    def fetch_seo_keywords(api_url):
        if not seo_keywords and not seo_description:
            st.error("Please provide at least one input (keywords or description).")
            return
        
        if not location_ids:
            st.error("Please select at least one location.")
            return
            
        if language_id is None:
            st.error("Please select a language.")
            return

        # Create payload with the new required parameters
        payload = {
            "keywords": seo_keywords, 
            "description": seo_description,
            "exclude_values": exclude_values,  
            "branded_keyword":st.session_state.seo_word_list, 
            "location_ids": location_ids,
            "language_id": language_id,
            "branded_words": remove_branded
        }
        
        with st.spinner("Generating keywords..."):
            response = requests.post(api_url, json=payload)

        if response.status_code == 200:
            data = response.json()
            st.success("SEO Keywords Generated Successfully!")

            # Display extracted keywords as JSON
            # st.subheader("Extracted SEO Keywords:")
            # st.json(data)

            # Convert JSON list to DataFrame
            try: 
                keywords_list = json.loads(data) if isinstance(data, str) else data
                if isinstance(keywords_list, list) and all(isinstance(item, dict) for item in keywords_list):
                    df = pd.DataFrame(keywords_list)
                    st.session_state.seo_df = df
                else:
                    st.error("Unexpected JSON format received!")
            except json.JSONDecodeError:
                st.error("Failed to parse JSON response!")
        else:
            error_msg = "Unknown error"
            try:
                error_detail = response.json().get('detail', error_msg)
                error_msg = error_detail
            except:
                pass
            st.error(f"Error: {error_msg}")

    

    # Function to process SEO keywords using the clustering API
    def process_seo_keywords():
        if st.session_state.seo_df is None or st.session_state.seo_df.empty:
            st.error("No SEO data available to process.")
            return

        # Convert DataFrame to CSV (in-memory)
        csv_buffer = BytesIO()
        st.session_state.seo_df.to_csv(csv_buffer, index=False)
        csv_buffer.seek(0)

        # Send CSV file to FastAPI
        files = {"file": ("seo_keywords.csv", csv_buffer, "text/csv")}
        response = requests.post(CLUSTER_API_URL, files=files)

        if response.status_code == 200:
            st.success("SEO Keywords Processed Successfully!")
            processed_data = response.json()

            # Convert JSON response to DataFrame
            try:
                processed_df = pd.DataFrame(processed_data)
                st.session_state.seo_processed_df = processed_df
                st.subheader("Processed SEO Keywords DataFrame:")
                st.dataframe(processed_df)
                return processed_df 
            except Exception as e:
                st.error(f"Error converting processed data to DataFrame: {e}")
                return None
        else:
            st.error(f"Error: {response.json().get('detail', 'Unknown error')}")
            return None

    # Buttons for generating and suggesting SEO keywords
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate SEO Keywords", key="gen_seo"):
            fetch_seo_keywords(GENERATE_API_URL)
            
    

    # Ensure session state variable exists
    if "seo_df" not in st.session_state:
        st.session_state.seo_df = pd.DataFrame({"Keyword": [], "Search Volume": []})

    # Show SEO DataFrame if available
    if st.session_state.seo_df is not None and not st.session_state.seo_df.empty:


        st.subheader("Editable SEO Keywords DataFrame:")

        # Add a checkbox column for deletion
        df = st.session_state.seo_df.copy()
        df["Delete"] = False  # Default unchecked
        edited_df = st.data_editor(df, key="seo_data_editor", num_rows="dynamic", use_container_width=True)

        # min_search = st.slider(
        # "Exclude Keywords Below Monthly Searches",
        # min_value=0,
        # max_value=int(st.session_state.seo_df["Avg Monthly Searches"].max()) if not st.session_state.seo_df.empty else 500,
        # value=0,
        # step=10,
        # key="seo_exclude"
        # )

        # # Apply the filter
        # filtered_df = st.session_state.seo_df[st.session_state.seo_df["Avg Monthly Searches"] >= min_search].copy()
        # Create button columns
        col1, col2, col3,col4, col5 = st.columns([1, 1, 1,1 ,1])  # Adjust layout as needed

        with col1:
            if st.button("Delete Selected Rows", key="delete_seo_rows"):
                st.session_state.seo_df = edited_df[edited_df["Delete"] == False].drop(columns=["Delete"])
                st.success("Selected rows deleted!")

        with col2:
            if st.button("Save SEO Changes", key="save_seo"):
                st.session_state.seo_df = edited_df.drop(columns=["Delete"])
                st.success("SEO Changes Saved!")

        with col5:
            if st.button("Process SEO Keywords", key="process_seo"):
                processed_df = process_seo_keywords()
                if processed_df is not None:
                    st.session_state.seo_processed_df = processed_df

                if "seo_processed_df" in st.session_state and not st.session_state.seo_processed_df is not None:
                    file_name = st.text_input("Enter filename (without extension):", "processed_seo")
                    if st.button("Upload to S3", key="upload_s3"):

                        csv_buffer = BytesIO()
                        st.session_state.seo_processed_df.to_csv(csv_buffer, index=False)
                        csv_buffer.seek(0)

                        # Ensure filename is valid
                        if not file_name.strip():
                            st.error("Filename cannot be empty!")
                        else:
                            full_filename = f"{file_name.strip()}.csv"
                            files = {"file": (full_filename, csv_buffer, "text/csv")}
                            response = requests.post(SEO_PROCESS_FILE, files=files)
            
                            if response.status_code == 200:
                                st.success(f"File uploaded to S3: {response.json()['s3_path']}")
                            
                            else:
                                st.error(f"Upload failed: {response.json().get('detail', 'Unknown error')}")

                        st.success("Uploading to S3...")  # Replace with actual S3 upload function
                   
                    if st.button("delete process csv", key="remove_upload_s3"):
                        del st.session_state["seo_processed_df"]
                        st.experimental_rerun()  # Refresh UI to remove the button
                else:
                    st.error("No processed SEO data available for upload.")     

        # Create another row of buttons
        # col4, col5 = st.columns([1, 1])

        with col4:
            if st.session_state.seo_df is not None and not st.session_state.seo_df.empty:
                csv = st.session_state.seo_df.to_csv(index=False)
                st.download_button(label="Download SEO CSV", data=csv, file_name="seo_keywords.csv", mime="text/csv", key="download_seo")

        with col3:
            if st.button("Clear SEO DataFrame", key="clear_seo"):
                st.session_state.seo_df = None
                st.session_state.seo_processed_df = None
                st.success("SEO Data Cleared!")
                
        # Show Processed SEO Keywords DataFrame if available
        if st.session_state.seo_processed_df is not None:
            st.subheader("Processed SEO Keywords DataFrame:")
            st.dataframe(st.session_state.seo_processed_df)

    # if "seo_processed_df" in st.session_state and not st.session_state.seo_processed_df is not None:
    #     if st.button("Upload to S3", key="upload_s3"):
    #         st.success("Uploading to S3...")  # Replace with actual S3 upload function
    #     # Add button to delete "Upload to S3" button from session
    #     if st.button("Remove Upload to S3 Button", key="remove_upload_s3"):
    #         del st.session_state["seo_processed_df"]
    #         st.experimental_rerun()  # Refresh UI to remove the button
    # else:
    #     st.error("No processed SEO data available for upload.")      

# PPC Tab
with tab2:
    st.header("PPC Keyword Generator")
    
    # User Input for PPC
    ppc_keywords = st.text_input("Enter PPC Keywords (comma-separated):", key="ppc_keywords", 
                                 placeholder="e.g., Google Ads, PPC Campaign, CPC")
    
    # Process input when user submits
    if ppc_keywords:
        keywords_list = [kw.strip() for kw in ppc_keywords.split(",") if kw.strip()]     
        keywords_list = keywords_list[:10]

        st.success(f"✅ Using the first {len(keywords_list)} keywords: {', '.join(keywords_list)}")


    ppc_description = st.text_area("Enter a Short Description (Optional):", key="ppc_description", 
                                   placeholder="Describe your PPC campaign goals")

    # Location and Language dropdowns
    st.subheader("Search Parameters")
    col1, col2 = st.columns(2)
    
    with col1:
        # Multi-select for locations
        selected_locations = st.multiselect(
            "Select Locations",
            options=[loc["country"] for loc in location_options],
            default=[location_options[0]["country"]],
            key="ppc_locations"
        )
        
        # Get the location IDs based on selected location names
        location_ids = [
            loc["id"] for loc in location_options 
            if loc["country"] in selected_locations
        ]
    
    with col2:
        # Single select for language
        selected_language = st.selectbox(
            "Select Language",
            options=[lang["Name"] for lang in language_options],
            index=0,
            key="ppc_language"
        )
        
        # Get the language ID based on selected language name
        language_id = next(
            (lang["ID"] for lang in language_options if lang["Name"] == selected_language),
            None
        )

    # Checkbox to enable/disable the exclusion filter
    use_exclude_filter = st.checkbox("Enable Exclusion Filter ppc")
    if use_exclude_filter:
        max_exclude = st.slider(
            "Exclude Keywords with Monthly Searches Below",
            min_value=0,
            max_value=500,
            value=0,
            step=10,
            key="ppc_exclude"
        )
    else:
        max_exclude = None   


    # Generate exclusion list
    if max_exclude is None:
        exclude_values = []
    else:
        exclude_values = list(range(0, max_exclude + 1, 10))  # Include max_exclude in the range

        # Ensure the max_exclude is in the list if it's not already there
        if max_exclude > 0 and max_exclude not in exclude_values:
            exclude_values.append(int(max_exclude))



    st.caption(f"Will exclude keywords with monthly searches in: {exclude_values}")



    if 'ppc_word_list' not in st.session_state:
        st.session_state.ppc_word_list = []

    ppc_new_word = st.text_input("Enter branded Keywords (comma-separated):", key="ppc_branded_key", 
                                    placeholder="e.g., apple, samsung, Automation")


    if st.button("Add branded Keywords"):
        if ppc_new_word:
            words = [word.strip() for word in ppc_new_word.split(',')]
    
            for word in words:
                if word and word not in st.session_state.ppc_word_list:
                    st.session_state.ppc_word_list.append(word)
            
      
    show_keywords = st.checkbox("Show branded Keywords", value=False)

    if show_keywords:
        st.subheader("Current Word List:")
        if st.session_state.ppc_word_list:
            # Create columns for each word to add a remove button
            for i, word in enumerate(st.session_state.ppc_word_list):
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.write(f"{i+1}. {word}")
                with col2:
                    # Create a unique key for each button
                    if st.button("Remove", key=f"remove_{i}"):
                        st.session_state.ppc_word_list.pop(i)
                        st.rerun()  # Rerun the app to update the display
        else:
            st.write("No words added yet.")


    
    def fetch_ppc_keywords(api_url):
        if not ppc_keywords and not ppc_description:
            st.error("Please provide at least one input (keywords or description).")
            return
        
        if not location_ids:
            st.error("Please select at least one location.")
            return
            
        if language_id is None:
            st.error("Please select a language.")
            return

    
        payload = {
            "keywords": ppc_keywords, 
            "description": ppc_description,
            "exclude_values": exclude_values,
            "branded_keyword":st.session_state.ppc_word_list,  
            "location_ids": location_ids,
            "language_id": language_id
        }
        
        with st.spinner("Generating keywords..."):
            response = requests.post(api_url, json=payload)

        if response.status_code == 200:
            data = response.json()
            st.success("PPC Keywords Generated Successfully!")

            # Display extracted keywords as JSON
            # st.subheader("Extracted PPC Keywords:")
            # st.json(data)

            # Convert JSON list to DataFrame
            try: 
                keywords_list = json.loads(data) if isinstance(data, str) else data
                if isinstance(keywords_list, list) and all(isinstance(item, dict) for item in keywords_list):
                    df = pd.DataFrame(keywords_list)
                    st.session_state.ppc_df = df
                else:
                    st.error("Unexpected JSON format received!")
            except json.JSONDecodeError:
                st.error("Failed to parse JSON response!")
        else:
            error_msg = "Unknown error"
            try:
                error_detail = response.json().get('detail', error_msg)
                error_msg = error_detail
            except:
                pass
            st.error(f"Error: {error_msg}")

    

    # Function to process SEO keywords using the clustering API
    def process_ppc_keywords():
        if st.session_state.ppc_df is None or st.session_state.ppc_df.empty:
            st.error("No PPC data available to process.")
            return

        # Convert DataFrame to CSV (in-memory)
        csv_buffer = BytesIO()
        st.session_state.ppc_df.to_csv(csv_buffer, index=False)
        csv_buffer.seek(0)

        # Send CSV file to FastAPI
        files = {"file": ("ppc_keywords.csv", csv_buffer, "text/csv")}
        response = requests.post(PPC_CLUSTER_API_URL, files=files)

        if response.status_code == 200:
            st.success("ppc Keywords Processed Successfully!")
            processed_data = response.json()

            # Convert JSON response to DataFrame
            try:
                processed_df = pd.DataFrame(processed_data)
                st.session_state.ppc_processed_df = processed_df
                # st.subheader("Processed SEO Keywords DataFrame:")
                # st.dataframe(processed_df)
            except Exception as e:
                st.error(f"Error converting processed data to DataFrame: {e}")
        else:
            st.error(f"Error: {response.json().get('detail', 'Unknown error')}")

    # Buttons for generating and suggesting SEO keywords
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate PPC Keywords", key="gen_ppc"):
            fetch_ppc_keywords(PPC_GENERATE_API_URL)
            
    

    # Show SEO DataFrame if available
    if st.session_state.ppc_df is not None and not st.session_state.ppc_df.empty:
        st.subheader("Editable PPC Keywords DataFrame:")

        # Add a checkbox column for deletion
        df = st.session_state.ppc_df.copy()
        df["Delete"] = False  # Default unchecked
        edited_df = st.data_editor(df, key="ppc_data_editor", num_rows="dynamic", use_container_width=True)

        col1, col2, col3,col4, col5 = st.columns([1, 1, 1, 1, 1])

        # Handle row deletion
        with col1:
            if st.button("Delete Selected Rows", key="delete_ppc_rows"):
                st.session_state.ppc_df = edited_df[edited_df["Delete"] == False].drop(columns=["Delete"])
                st.success("Selected rows deleted!")

        # Save updated DataFrame
        with col2:
            if st.button("Save PPC Changes", key="save_ppc"):
                st.session_state.ppc_df = edited_df.drop(columns=["Delete"])
                st.success("PPC Changes Saved!")

        # Download the modified DataFrame as CSV
        with col3:
            if st.session_state.ppc_df is not None and not st.session_state.ppc_df.empty:
                csv = st.session_state.ppc_df.to_csv(index=False)
                st.download_button(label="Download PPC CSV", data=csv, file_name="ppc_keywords.csv", mime="text/csv", key="download_ppc")

        # Process SEO keywords (send to clustering API)
        with col5:
            if st.button("Process PPC Keywords", key="process_ppc"):
                process_ppc_keywords()

        # Clear SEO DataFrame button
        with col4:
            if st.button("Clear PPC DataFrame", key="clear_ppc"):
                st.session_state.ppc_df = None
                st.session_state.ppc_processed_df = None
                st.success("PPC Data Cleared!")
            
    # Show Processed SEO Keywords DataFrame if available
    if st.session_state.ppc_processed_df is not None:
        st.subheader("Processed SEO Keywords DataFrame:")
        st.dataframe(st.session_state.ppc_processed_df)



with tab3:
    st.header("Keywords Suggestion")
    
    # User Input for SEO
    seo_keywords = st.text_input("Enter SEO Keywords (comma-separated):", key="suggestion_key", 
                                 placeholder="e.g., SEO, Machine Learning, Automation")
    seo_description = st.text_area("Enter a Short Description (Optional):", key="suggestion_desc", 
                                   placeholder="Describe your business or service for SEO")        
    

    # Function to fetch additional suggested SEO keywords
    def fetch_suggested_seo_keywords():
        if not seo_keywords:
            st.error("Please provide some keywords for suggestions.")
            return

        payload = {"keywords": seo_keywords, "description": seo_description}
        response = requests.post(SUGGEST_API_URL, json=payload)
        
        if response.status_code == 200:
            suggested_keywords = response.json()
            st.success("Suggested SEO Keywords Generated!")

            st.write(suggested_keywords)
        else:
            st.error(f"Error: {response.json().get('detail', 'Unknown error')}")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Suggest More SEO Keywords", key="suggest_seo"):
            fetch_suggested_seo_keywords()        


def get_documents(user_id: str = "User", category: Optional[str] = None) -> List[str]:
    """Fetch a list of documents from the API."""
    # print(f"{FETCH_DOCUMENTS_URL}/{user_id}/{category}")
    response = requests.get(f"{FETCH_DOCUMENTS_URL}/{user_id}/{category}")
    # print(response)
    if response.status_code == 200:
        return response.json().get("documents", [])
    return []


with tab4:
    # Streamlit app title
    # Organize documents into folders
    folders = {
        "Buyer persona": [],
        "Tone of voice": [],
        "Brand identity": [],
        "Offering": []
    }
    
    user_id = "User"  
    for category in folders.keys():
        folders[category] = get_documents(user_id, category)
        # print(folders[category])s
    
    print(f"folder",folders)
    st.subheader("Select a Document from S3")
    # Dropdowns for each folder
    selected_documents = {}
    for folder, docs in folders.items():
        selected_documents[folder] = st.selectbox(
            f"Select a document from {folder}", 
            options=docs, 
            key=folder
        )
        # print(docs)
        # print(folder)
        # print(selected_documents[folder])
    print(selected_documents)

    if st.button("🔄 Reload", key="reload_button"):
        st.rerun()
    # Display download buttons
    # for folder, doc in selected_documents.items():
    #     if doc != "None":
    #         if st.button(f"Download {folder} Document"):
    #             file_url = fetch_document(doc)
    #             if file_url:
    #                 st.markdown(f"[Download {doc.split('/')[-1]}]({file_url})")
    #             else:
    #                 st.error("Error fetching the document.")

    st.subheader("Social Media Post (campaign)")

    uploaded_file = st.file_uploader("Upload a Word document", type=["docx", "doc"])

    if uploaded_file is not None:

        if st.button("Submit File"):

            try:
                payload = {
                    "data": selected_documents
                }

                response = requests.post(DOWNLOAD_DOCUMENT_URL, json=payload)
                print(response)
                if response.status_code == 200:
                    # extracted_texts = response.json().get("extracted_texts", {})
                    summarized_data = response.json()
                    # st.json(response.json())
                    # st.success("File uploaded successfully!")
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")    
                
            # Prepare the file for upload
            # files = {"file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type)}
            files = {
                    "file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type),
                    "json_data": (None, json.dumps(summarized_data), "application/json")
                    }

            
            try:
                # json_payload = {"json_data": json.dumps(summerized_data)}
                # Make the API request
                response = requests.post(SOCIAL_MEDIA_API_URL, files=files)#,json= json_payload)
                
                # Check if the response was successful
                if response.status_code == 200:
                    st.write("### LinkedIn, Facebook, Twitter posts")
                    # st.json(response.json())  
                    data = response.json()
                    df_data = [
                                {
                                    "LinkedIn": iteration["LinkedIn"][0],
                                    "Facebook": iteration["Facebook"][0],
                                    "Twitter": iteration["Twitter"][0],
                                    "Image Headline": iteration["Image Headline"][0],
                                    "Subheadline": iteration["Subheadline"][0]
                                }
                                for iteration in data
                            ]

                    df = pd.DataFrame(df_data)
                    st.dataframe(df)
                else:
                    # Handle error cases without assuming JSON
                    st.error(f"Error: {response.status_code}, {response.text}")
                    
            except requests.exceptions.RequestException as e:
                st.error(f"Request failed: {str(e)}")       






def convert_doc_to_docx(input_file_path):
    """Convert .doc to .docx using LibreOffice (Linux/Mac) or pywin32 (Windows).
    
    Args:
        input_file_path (str): Path to the input .doc file
    
    Returns:
        str: Path to the converted .docx file
    """
    # Ensure we only replace the last .doc extension and avoid double extensions
    base_name, ext = os.path.splitext(input_file_path)
    if ext.lower() == '.doc':
        output_file = base_name + ".docx"
    else:
        output_file = input_file_path  # If not .doc, keep original (shouldn't happen here)
    
    try:
        if os.name == "nt":  # Windows
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(input_file_path))
            doc.SaveAs(os.path.abspath(output_file), 16)  # 16 = docx format
            doc.Close()
            word.Quit()
        else:  # Linux/Mac using LibreOffice
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                os.path.dirname(input_file_path) or ".",
                input_file_path
            ], check=True)
        
        return output_file
    
    except Exception as e:
        raise Exception(f"Conversion failed: {str(e)}")

def read_docx(file):
    """Read content from a .docx file.
    
    Args:
        file: File object or path to .docx file
    
    Returns:
        str: Text content of the document
    """
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text if text else "Document is empty"
    except Exception as e:
        raise Exception(f"Error reading document: {str(e)}")

def save_docx(text, filename="edited_document.docx"):
    """Save text as a .docx file.
    
    Args:
        text (str): Text content to save
        filename (str): Name for the output file
    
    Returns:
        BytesIO: Buffer containing the saved document
    """
    try:
        doc = Document()
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        raise Exception(f"Error saving document: {str(e)}")

           

with tab5:
    st.subheader("Upload & Edit Documents")

    categories = {
        "Buyer persona": "buyer",
        "Tone of voice": "tone",
        "Brand identity": "brand",
        "Offering": "offering"
    }

    for category, key in categories.items():
        st.write(f"### {category}")
        uploaded_file = st.file_uploader(
            f"Upload a document for {category}",
            type=["doc", "docx"],
            key=key,
            accept_multiple_files=False
        )

        if uploaded_file is not None:
            try:
                # Handle file based on extension
                file_extension = uploaded_file.name.lower().split('.')[-1]
                temp_file_path = None
                
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
                    temp_file.write(uploaded_file.getvalue())
                    temp_file_path = temp_file.name

                # Convert .doc to .docx if necessary
                if file_extension == "doc":
                    converted_path = convert_doc_to_docx(temp_file_path)
                    file_to_read = converted_path
                else:
                    file_to_read = temp_file_path

                # Read and display document content
                file_text = read_docx(file_to_read)
                edited_text = st.text_area(
                    f"Edit the document ({category})",
                    file_text,
                    height=300,
                    key=f"edit_{key}"
                )

                # Prepare files and data for upload
                files = {"file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type)}
                data = {"category": category}

                # Combined Upload/Save/Edit button
                if st.button(f"Save {category}", key=f"upload_save_{key}"):
                    # If text was edited, save it first
                    if edited_text != file_text:
                        output_filename = os.path.splitext(uploaded_file.name)[0] + ".docx"
                        updated_doc = save_docx(edited_text)
                        files = {
                            "file": (output_filename, updated_doc, 
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        }
                        st.success(f"Changes saved for {category}")
                        
                        # Provide download button for edited file
                        st.download_button(
                            label=f"Download Edited {category} Document",
                            data=updated_doc,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{key}"
                        )

                if st.button(f"Upload {category}", key=f"upload_{key}"):
                    response = requests.post(UPLOAD_FILE_S3_BUCKET, files=files, data=data)
                    if response.status_code == 200:
                        st.success(f"File uploaded to S3: {response.json()['s3_path']}")
                        st.write(f"Uploaded file: {uploaded_file.name}")
                    else:
                        st.error(f"Upload failed: {response.json().get('detail', 'Unknown error')}")

                # Clean up temporary files
                if temp_file_path and os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                if file_extension == "doc" and os.path.exists(converted_path):
                    os.unlink(converted_path)

            except Exception as e:
                st.error(f"Error processing {category} document: {str(e)}")
                # Clean up on error
                if temp_file_path and os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                if 'converted_path' in locals() and os.path.exists(converted_path):
                    os.unlink(converted_path)


with tab6:
    st.subheader("SEO Content generation")

    folders_cg = {
        "Buyer persona": [],
        "Tone of voice": [],
        "Brand identity": [],
        "Offering": []
    }
    
    user_id = "User"  
    for category in folders_cg.keys():
        folders_cg[category] = get_documents(user_id, category)
        # print(folders[category])
    

    st.write("Select a Document from S3")
    # Dropdowns for each folder
    selected_documents_cg = {}
    for folder, docs in folders_cg.items():
        selected_documents_cg[folder] = st.selectbox(
            f"Select a documents from {folder}", 
            options=docs, 
            key=f"{folder}_ss"
        )
    
    print(selected_documents_cg)

        
    upload_seo_file = st.checkbox("Upload keywords file", value=False)
    if upload_seo_file:
        uploaded_file = st.file_uploader("Choose a CSV file", type=["csv"])
        if uploaded_file:
            if st.button("Upload to S3"):
                files = {"file": (uploaded_file.name, uploaded_file.getvalue(), "text/csv")}
                response = requests.post(SEO_PROCESS_FILE, files=files)

                if response.status_code == 200:
                    s3_path = response.json()["s3_path"]
                    st.success(f"File uploaded successfully! 🚀")
                    st.write(f"**S3 Path:** {s3_path}")
                else:
                    st.error(f"Upload failed: {response.json().get('detail', 'Unknown error')}")


    content_folder = {"csv": []}

    user_id = "User"  
    folder_seo = "seo_content_generation"

    # Fetch documents for each category
    for category in content_folder.keys():
        content_folder[category] = get_documents(user_id, folder_seo)

    # st.write("Select a Document from S3")

    # Dropdowns for each folder category
    selected_csv = {}
    for folder, docs in content_folder.items():
        if docs:  # Show dropdown only if there are documents
            selected_csv[folder] = st.selectbox(
                f"Select a document from {folder}",
                options=docs,
                key=folder
            )
        else:
            st.warning(f"No documents found in {folder}")

    if st.button("🔄 Reload", key="reload_button_seo"):
        st.rerun()        
    
    st.subheader("Blog content generation")

    uploaded_file_bg = st.file_uploader("Upload a blog content", type=["docx", "doc"])

    if uploaded_file_bg is not None:
        if st.button("Submit File"):
            try:
                    
                # payload = {
                #         "data": selected_csv
                #     }
                # print(payload)
                # response_csv = requests.post(DOWNLOAD_CONTENT_CSV, json=payload)
                # if response_csv.status_code == 200:
                #     json_data = response_csv.json()
                #     print(json_data)
                payload = {
                        "data": selected_documents_cg
                    }

                response_summ = requests.post(DOWNLOAD_DOCUMENT_URL, json=payload)
                # print(response)
                if response_summ.status_code == 200:
                    summarized_data = response_summ.json()
                    print(summarized_data)


                files = {
                    "file": (uploaded_file_bg.name, uploaded_file_bg.getvalue(), uploaded_file_bg.type),
                    "json_data": (None, json.dumps(summarized_data), "application/json")
                    }

                try:
                    response = requests.post(BLOG_GENERATION_TEXT, files=files)#,json= json_payload)
                    if response.status_code == 200:
                       data = response.json()
                       st.session_state.blog_text = data
                    else:
                        st.error(f"Error: {response.status_code}, {response.text}")
                        
                except requests.exceptions.RequestException as e:
                    st.error(f"Request failed: {str(e)}") 
                
                # Display the editor
                st.title("Edit Your Blog Post")
                edited_text = st.text_area(
                    "Blog Post Content",
                    value=st.session_state.blog_text,
                    height=600,  # Tall enough for the full blog
                    key="blog_editor"
                )
                st.session_state.blog_text = edited_text

            except requests.exceptions.RequestException as e:
                st.error(f"Request failed: {str(e)}")       

    

   

