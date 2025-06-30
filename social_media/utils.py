from concurrent.futures import ThreadPoolExecutor
import os
import pypandoc
from docx import Document
from io import BytesIO
from docx import Document
import tempfile
from auth.auth import get_db
from auth.models import SocialMediaFile
from datetime import datetime, timedelta
from typing import Union
# def extract_text_from_docx(file_path):
#     """Extracts text from a .docx file."""
#     doc = Document(file_path)
#     return '\n'.join([para.text for para in doc.paragraphs])

# def extract_text_from_doc(file_path):
#     """Extracts text from a .doc file using Pandoc (cross-platform)."""
#     return pypandoc.convert_file(file_path, 'plain')

# def convert_doc_to_text(file_path):
#     """Converts a .doc or .docx file into text."""
#     if not os.path.exists(file_path):
#         raise FileNotFoundError("File not found.")
    
#     file_ext = os.path.splitext(file_path)[1].lower()
#     if file_ext == '.docx':
#         return extract_text_from_docx(file_path)
#     elif file_ext == '.doc':
#         return extract_text_from_doc(file_path)
#     else:
#         raise ValueError("Unsupported file format. Use .doc or .docx")


def extract_text_from_docx_bytes(file_bytes):
    """Extracts text from a .docx file in bytes."""
    doc = Document(BytesIO(file_bytes))
    return '\n'.join([para.text for para in doc.paragraphs])

def extract_text_from_doc_bytes(file_bytes):
    """Extracts text from a .doc file in bytes using Pandoc."""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
        temp_file.write(file_bytes)
        temp_file_path = temp_file.name
    try:
        return pypandoc.convert_file(temp_file_path, 'plain')
    finally:
        os.unlink(temp_file_path)

def convert_doc_to_text(file_bytes, filename):
    """Converts a .doc or .docx file in bytes into text."""
    file_ext = os.path.splitext(filename)[1].lower()
    if file_ext == '.docx':
        return extract_text_from_docx_bytes(file_bytes)
    elif file_ext == '.doc':
        return extract_text_from_doc_bytes(file_bytes)
    else:
        raise ValueError("Unsupported file format. Use .doc or .docx")
# Example usage
# text = convert_doc_to_text(r"C:\Users\nickc\OneDrive\Desktop\AI marketing\doc\Plug_%26_Play%2C_Grow_-_campaing (1).docx")
# print(text)

import re
import emoji

def clean_post(text, remove_emojis=True, remove_hashtags=True):
    if remove_emojis == False:
        text = ''.join(char for char in text if not emoji.is_emoji(char))

    if remove_hashtags == False:
        text = re.sub(r'#\w+', '', text)

    return text.strip()


def clean_post_list(data_list, remove_emojis=True, remove_hashtags=True): 
    cleaned = []
    for item in data_list:
        new_item = {}
        for key, val in item.items():
            if isinstance(val, list) and all(isinstance(v, str) for v in val):
                new_item[key] = [
                    clean_post(
                        text,
                        remove_emojis=remove_emojis,
                        remove_hashtags=remove_hashtags,
                    ) for text in val
                ]
            else:
                new_item[key] = val  # leave non-list or non-str values untouched
        cleaned.append(new_item)
    return cleaned


def upload_socialmedia_table( uuid: str, user_id: int, file_name: str, linkedIn: Union[dict, list], facebook_post: Union[dict, list], twitter_post: Union[dict, list]):
    db = next(get_db()) 
    try:
        new_file = SocialMediaFile(
            user_id=user_id,
            file_name=file_name,
            uuid=uuid,
            linkedIn_post = linkedIn,
            facebook_post = facebook_post,
            twitter_post  = twitter_post,
            last_reset = datetime.utcnow()
        )

        
        db.add(new_file)
        db.commit()
        db.refresh(new_file)

        return {
            "id": new_file.id,  # Assuming an id field exists
            "user_id": new_file.user_id,
            "file_name": new_file.file_name,
            "uuid": new_file.uuid,
            "linkedIn_post": new_file.linkedIn_post,
            "facebook_post": new_file.facebook_post,
            "twitter_post": new_file.twitter_post,
            "last_reset": new_file.last_reset.isoformat() if new_file.last_reset else None
        }  # You can return the created object if needed

    except Exception as e:
        db.rollback()
        print(str(e))
        raise Exception(f"Error storing SEO file in table: {str(e)}")
    finally:
        db.close() 


# Optimized version with key improvements

# @router.post("/social_media_post")
# async def social_media_post(
#     file: Optional[UploadFile] = File(None),
#     fileName: Optional[str] = Form(None),
#     text_data: Optional[str] = Form(None),
#     json_data: Optional[str] = Form(None),
#     linkedIn_post: Optional[bool] = Form(True),
#     facebook_post: Optional[bool] = Form(True),
#     twitter_post: Optional[bool] = Form(True),
#     hash_tag: Optional[bool] = Form(False),
#     emoji: Optional[bool] = Form(False),
#     objectives: Optional[str] = Form(None),
#     audience: Optional[str] = Form(None),
#     user=Depends(check_api_limit("social_media")),
#     db: Session = Depends(get_db),
#     id: str = Depends(verify_jwt_token)
# ):
#     user_id = int(id[1])
#     unique_id = uuid.uuid4().hex

#     try:
#         # 1. Early validation
#         await validate_input(file, text_data)
        
#         # 2. Parse JSON data once
#         parsed_objectives = parse_json_safely(objectives)
#         parsed_audience = parse_json_safely(audience)
        
#         # 3. Optimized database query with single query + caching
#         summarized_data = await get_summarized_data(
#             db, user_id, parsed_objectives, parsed_audience
#         )
        
#         # 4. Async file processing
#         text = await process_file_content(file, text_data)
        
#         # 5. Optimized parallel execution
#         results, total_tokens = await execute_social_media_tasks(
#             text, summarized_data, linkedIn_post, facebook_post, 
#             twitter_post, hash_tag, emoji
#         )
        
#         # 6. Batch database operations
#         await update_database_records(
#             db, user, user_id, unique_id, fileName, 
#             results, total_tokens
#         )
        
#         return {
#             "uuid": unique_id,
#             "fileName": fileName,
#             "data": results
#         }
        
#     except ValidationError as e:
#         await handle_error_cleanup(db, user)
#         raise HTTPException(status_code=400, detail=str(e))
#     except Exception as e:
#         await handle_error_cleanup(db, user)
#         logger.error(f"Social media post error: {str(e)}", exc_info=True)
#         raise HTTPException(status_code=500, detail="Internal server error")


# Helper functions for better separation of concerns
# from typing import Optional
# from fastapi import APIRouter, UploadFile

# import asyncio
# from functools import lru_cache
# import json
# from sqlalchemy.orm import Session
# import asyncio
# from sqlalchemy.orm.attributes import flag_modified
# from concurrent.futures import ThreadPoolExecutor
# from social_media.Agents.document_summared import Document_summerizer


# async def validate_input(file: Optional[UploadFile], text_data: Optional[str]):
#     """Validate input parameters early"""
#     if not file and not text_data:
#         raise ValidationError("Either file or text_data must be provided")
    
#     if file and not file.filename.lower().endswith((".docx", ".doc")):
#         raise ValidationError("Invalid file format. Please upload a .docx or .doc file")


# def parse_json_safely(json_str: Optional[str]) -> list:
#     """Safely parse JSON string with error handling"""
#     if not json_str:
#         return []
#     try:
#         return json.loads(json_str)
#     except json.JSONDecodeError:
#         logger.warning(f"Failed to parse JSON: {json_str}")
#         return []


# @lru_cache(maxsize=128)
# async def get_cached_file_content(user_id: int, uuid_ids: tuple) -> dict:
#     """Cache frequently accessed file content"""
#     # Implementation for caching logic
#     pass


# async def get_summarized_data(
#     db: Session, user_id: int, objectives: list, audience: list
# ) -> dict:
#     """Optimized data retrieval with single query and caching"""
#     if not objectives and not audience:
#         return {}
    
#     # Single query for all UUIDs
#     all_uuids = set(objectives + audience)
#     if not all_uuids:
#         return {}
    
#     # Use async query if available, or execute in thread pool
#     file_contents = await asyncio.get_event_loop().run_in_executor(
#         None, 
#         lambda: db.query(SourceFileContent)
#         .filter(
#             SourceFileContent.user_id == user_id,
#             SourceFileContent.uuid_id.in_(all_uuids)
#         )
#         .all()
#     )
    
#     if not file_contents:
#         return {}
    
#     # Create lookup dict for O(1) access
#     content_lookup = {fc.uuid_id: fc.file_data for fc in file_contents}
    
#     # Collect relevant data
#     file_data_list = []
#     for uuid_id in all_uuids:
#         if uuid_id in content_lookup:
#             file_data_list.append(content_lookup[uuid_id])
    
#     if file_data_list:
#         summarized_data, _ = await asyncio.get_event_loop().run_in_executor(
#             None, Document_summerizer, file_data_list
#         )
#         return summarized_data
    
#     return {}


# async def process_file_content(file: Optional[UploadFile], text_data: Optional[str]) -> str:
#     """Process file content asynchronously"""
#     if file:
#         file_contents = await file.read()
#         return await asyncio.get_event_loop().run_in_executor(
#             None, convert_doc_to_text, file_contents, file.filename
#         )
#     else:
#         return text_data


# async def execute_social_media_tasks(
#     text: str, summarized_data: dict, linkedIn_post: bool,
#     facebook_post: bool, twitter_post: bool, hash_tag: bool, emoji: bool
# ) -> tuple[dict, int]:
#     """Execute social media tasks with better error handling and resource management"""
    
#     # Use ThreadPoolExecutor with limited workers to prevent resource exhaustion
#     max_workers = min(3, os.cpu_count() or 1)
    
#     with ThreadPoolExecutor(max_workers=max_workers) as executor:
#         loop = asyncio.get_event_loop()
#         tasks = []
#         task_types = []
        
#         if linkedIn_post:
#             tasks.append(
#                 loop.run_in_executor(
#                     executor, linkedIn_agent_call, text, summarized_data, 5, hash_tag, emoji
#                 )
#             )
#             task_types.append("linkedin")
        
#         if facebook_post:
#             tasks.append(
#                 loop.run_in_executor(
#                     executor, facebook_agent_call, text, summarized_data, 5, hash_tag, emoji
#                 )
#             )
#             task_types.append("facebook")
        
#         if twitter_post:
#             tasks.append(
#                 loop.run_in_executor(
#                     executor, twitter_agent_call, text, summarized_data, 5, hash_tag, emoji
#                 )
#             )
#             task_types.append("twitter")
        
#         if not tasks:
#             return {}, 0
        
#         # Execute with timeout to prevent hanging
#         try:
#             responses = await asyncio.wait_for(
#                 asyncio.gather(*tasks, return_exceptions=True), 
#                 timeout=300  # 5 minute timeout
#             )
#         except asyncio.TimeoutError:
#             logger.error("Social media tasks timed out")
#             raise Exception("Request timed out")
        
#         # Process results with error handling
#         results = {}
#         total_tokens = 0
        
#         for i, (response, task_type) in enumerate(zip(responses, task_types)):
#             if isinstance(response, Exception):
#                 logger.error(f"{task_type} task failed: {str(response)}")
#                 continue
                
#             try:
#                 data, tokens = response
#                 results[f"{task_type}_posts"] = data
#                 total_tokens += tokens
#             except (ValueError, TypeError) as e:
#                 logger.error(f"Failed to process {task_type} response: {str(e)}")
#                 continue
        
#         return results, total_tokens


# async def update_database_records(
#     db: Session, user, user_id: int, unique_id: str, 
#     fileName: str, results: dict, total_tokens: int
# ):
#     """Batch database operations for better performance"""
    
#     def db_operations():
#         try:
#             # Update social media record
#             social_media_record = db.query(SocialMedia).filter(
#                 SocialMedia.user_id == user.id
#             ).first()
            
#             if social_media_record:
#                 if total_tokens > 0:
#                     social_media_record.total_tokens += total_tokens
#                 if not results:
#                     social_media_record.call_count = max(
#                         social_media_record.call_count - 1, 0
#                     )
            
#             # Upload social media data
#             if results:
#                 linkedin_file = results.get("linkedin_posts", [])
#                 facebook_file = results.get("facebook_posts", [])
#                 twitter_file = results.get("twitter_posts", [])
                
#                 upload_socialmedia_table(
#                     str(unique_id), user_id, fileName,
#                     linkedIn=linkedin_file, 
#                     facebook_post=facebook_file, 
#                     twitter_post=twitter_file
#                 )
            
#             db.commit()
            
#         except Exception as e:
#             db.rollback()
#             raise e
    
#     # Execute database operations in thread pool
#     await asyncio.get_event_loop().run_in_executor(None, db_operations)


# async def handle_error_cleanup(db: Session, user):
#     """Handle error cleanup operations"""
#     def cleanup():
#         try:
#             social_media_record = db.query(SocialMedia).filter(
#                 SocialMedia.user_id == user.id
#             ).first()
#             if social_media_record:
#                 social_media_record.call_count = max(
#                     social_media_record.call_count - 1, 0
#                 )
#                 db.commit()
#         except Exception as e:
#             logger.error(f"Error during cleanup: {str(e)}")
#             db.rollback()
    
#     await asyncio.get_event_loop().run_in_executor(None, cleanup)


# # Additional optimizations to consider:

# class SocialMediaConfig:
#     """Configuration class for better maintainability"""
#     MAX_WORKERS = 3
#     TASK_TIMEOUT = 300
#     CACHE_SIZE = 128
#     ALLOWED_EXTENSIONS = {".docx", ".doc"}

# from pydantic import BaseModel
# # Response model for better API documentation
# class SocialMediaResponse(BaseModel):
#     uuid: str
#     fileName: Optional[str]
#     data: dict




# # Custom exceptions
# class ValidationError(Exception):
#     pass        