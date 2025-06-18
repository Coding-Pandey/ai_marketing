from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Body, Form, Query
from fastapi.responses import JSONResponse
from auth.models import SourceFileCategory, SourceFileContent
from auth.auth import get_db
from sqlalchemy.orm.attributes import flag_modified
from sqlalchemy.orm import Session
from utils import verify_jwt_token, check_api_limit
from docx import Document
from io import BytesIO
import uuid
import datetime
# import textract 

router = APIRouter()

@router.get("/uploaded_files")
async def get_uploaded_files(user_id: str = Depends(verify_jwt_token), db: Session = Depends(get_db)):
    user_id = user_id[1]
    files = db.query(SourceFileContent).filter(SourceFileContent.user_id == user_id).all()
    if not files:
        return []
    files = [
        {
            "file_name": file.file_name,
            "category": file.category,
            "uuid_id": file.uuid_id,
            "uploaded_file_name":file.extracted_text
        } for file in files
    ]

    return {"uploaded_files": files}

@router.delete("/delete_Source_file/{uuid_id}")
async def delete_source_file(uuid_id: str, user_id: str = Depends(verify_jwt_token), db: Session = Depends(get_db)):
    user_id = user_id[1]
    file_record = db.query(SourceFileContent).filter(
        SourceFileContent.uuid_id == uuid_id,
        SourceFileContent.user_id == user_id
    ).first()

    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")

    db.delete(file_record)
    db.commit()
    flag_modified(file_record, "file_data")  # Mark file_data as modified if needed

    return {"message": "File deleted successfully"}

@router.post("/upload_and_parse")
async def upload_and_parse_file(
    category: SourceFileCategory = Form(...),
    file:UploadFile = File(...),
    file_name: str = Form(...),
    user_id: str = Depends(verify_jwt_token),
    db: Session = Depends(get_db)
):
    # Extract text from file
    user_id = user_id[1]
    uploaded_file_name = file.filename

    file_ext = file.filename.lower().split(".")[-1]
    file_bytes = await file.read()

    if file_ext == "docx":
        
        doc = Document(BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        print(text)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")
    
    print(category, file.filename, user_id)
    json_data = {category.name: text}
    uuid_id = str(uuid.uuid4().hex)
    # Save text to DB
    record = SourceFileContent(
        user_id=user_id,
        uuid_id=uuid_id,
        file_name=file_name,
        category=category,
        uploaded_at=datetime.datetime.now(),
        extracted_text=uploaded_file_name,
        file_data = json_data
    )
    db.add(record)
    db.commit()
    db.refresh(record)

    return {"message": "File content saved", "record_id": record.id, "file_name":file_name, "uuid_id": uuid_id, "uploaded_file_name":uploaded_file_name}

@router.patch("/upload_and_parse/{uuid_id}")
async def upload_and_parse_file(
    uuid_id: str,
    category: SourceFileCategory = Form(...),
    file: UploadFile = File(None),
    file_name: str = Form(None),
    user_id: str = Depends(verify_jwt_token),
    db: Session = Depends(get_db)
):

    user_id = user_id[1]

    # Find existing record
    record = db.query(SourceFileContent).filter(
        SourceFileContent.uuid_id == uuid_id,
        SourceFileContent.user_id == user_id
    ).first()

    if not record:
        raise HTTPException(status_code=404, detail="File record not found")

    if file:
        file_ext = file.filename.lower().split(".")[-1]
        file_bytes = await file.read()

        if file_ext == "docx":
            uploaded_file_name = file.filename
            doc = Document(BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
        # elif file_ext == "doc":
        #     try:
        #         text = textract.process(BytesIO(file_bytes), extension='doc').decode('utf-8')
        #     except Exception as e:
        #         raise HTTPException(status_code=500, detail=f"Error processing .doc file: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Only .docx and .doc are supported.")

        # Update record
        json_data = {category.name: text}
        record.file_data = json_data
        record.extracted_text = uploaded_file_name

    if file_name:
        record.file_name = file_name

    flag_modified(record, "file_data")  # Mark the JSON column as modified
    db.commit()
    db.refresh(record)

    return {"message": "File content saved", "record_id": record.id,"file_name":file_name, "uuid_id": uuid_id}

# @router.get("/file_content")
# async def get_file_content(user_id: str = Depends(verify_jwt_token), db: Session = Depends(get_db)):
#     user_id = user_id[1]
#     records = db.query(SourceFileContent).filter(
#         SourceFileContent.user_id == user_id
#     ).all()

#     if not records:
#         raise HTTPException(status_code=404, detail="No file records found")

#     return [
#         {
#             "uuid_id": record.uuid_id,
#             "file_name": record.file_name,
#             "category": record.category,
#             "uploaded_file_name": record.extracted_text,
#             # "file_data": record.file_data
#         }
#         for record in records
#     ]

@router.get("/uploaded_sourcefiles")
async def get_uploaded_sourcefiles(user_id: str = Depends(verify_jwt_token), db: Session = Depends(get_db)):
    try:
        # Extract user_id
        user_id = user_id[1]
        # Query the database
        files = db.query(SourceFileContent).filter(SourceFileContent.user_id == user_id).all()
        # Log details of each file
        if not files:
            return {"define_objective": "No Post Objective uploaded.", "Target_audience": "No Audience uploaded."}

        result = {"define_objective": [], "Target_audience": []}
        
        for file in files:
            file_data = {
                "file_name": file.file_name,        
                "category": file.category if file.category else "Unknown",
                "uuid_id": file.uuid_id,
                "uploaded_file_name": file.extracted_text
            }
            
            if file.category.name == "BUYER_PERSONA":
                result["Target_audience"].append(file_data)
            else:
                result["define_objective"].append(file_data)
        
        return result
    
    except Exception as e:
    
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")
    





